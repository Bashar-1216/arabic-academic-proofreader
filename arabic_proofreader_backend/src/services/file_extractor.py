import os
import tempfile
import docx
import fitz  # PyMuPDF
from typing import Dict, Optional
import re

class FileExtractor:
    """Enhanced file extraction service for PDF and Word documents"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc'}
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.supported_extensions
    
    def validate_file_size(self, file_size: int) -> bool:
        """Validate file size"""
        return file_size <= self.max_file_size
    
    def extract_from_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF with enhanced error handling and metadata"""
        try:
            doc = fitz.open(file_path)
            
            # Extract metadata
            metadata = doc.metadata
            page_count = doc.page_count
            
            # Extract text from all pages
            full_text = ""
            page_texts = []
            
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                page_texts.append({
                    'page_number': page_num + 1,
                    'text': page_text,
                    'word_count': len(page_text.split())
                })
                full_text += page_text + "\n"
            
            doc.close()
            
            # Clean extracted text
            full_text = self._clean_extracted_text(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'creator': metadata.get('creator', ''),
                    'page_count': page_count,
                    'file_type': 'PDF'
                },
                'pages': page_texts,
                'stats': {
                    'total_words': len(full_text.split()),
                    'total_characters': len(full_text),
                    'pages': page_count
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'خطأ في قراءة ملف PDF: {str(e)}',
                'error_type': 'pdf_extraction_error'
            }
    
    def extract_from_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX with enhanced features"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:  # Skip empty paragraphs
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'text': para_text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
                    full_text += para_text + "\n"
            
            # Extract text from tables
            tables_text = []
            for table_num, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_content.append(row_content)
                
                tables_text.append({
                    'table_number': table_num + 1,
                    'content': table_content
                })
                
                # Add table text to full text
                for row in table_content:
                    full_text += " | ".join(row) + "\n"
            
            # Extract document properties
            core_props = doc.core_properties
            
            # Clean extracted text
            full_text = self._clean_extracted_text(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'file_type': 'DOCX'
                },
                'structure': {
                    'paragraphs': paragraphs,
                    'tables': tables_text
                },
                'stats': {
                    'total_words': len(full_text.split()),
                    'total_characters': len(full_text),
                    'paragraphs': len(paragraphs),
                    'tables': len(tables_text)
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'خطأ في قراءة ملف Word: {str(e)}',
                'error_type': 'docx_extraction_error'
            }
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newline
        text = re.sub(r'[ \t]+', ' ', text)      # Multiple spaces/tabs to single space
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\n\d+\n', '\n', text)   # Standalone page numbers
        text = re.sub(r'\n\d+\s*$', '', text, flags=re.MULTILINE)  # Page numbers at end of line
        
        # Clean up common PDF extraction artifacts
        text = re.sub(r'[^\w\s\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF.,;:!?()[\]{}"\'-]', '', text)
        
        # Normalize Arabic text
        text = self._normalize_arabic_text(text)
        
        return text.strip()
    
    def _normalize_arabic_text(self, text: str) -> str:
        """Normalize Arabic text"""
        # Normalize Arabic letters
        text = re.sub(r'[إأآا]', 'ا', text)  # Normalize Alif variations
        text = re.sub(r'[ىي]', 'ي', text)    # Normalize Ya variations
        text = re.sub(r'[ةه]', 'ة', text)    # Normalize Ta Marbuta
        
        # Remove diacritics (optional - might want to keep for academic texts)
        # text = re.sub(r'[\u064B-\u0652\u0670\u0640]', '', text)
        
        return text
    
    def extract_text(self, file_path: str, filename: str) -> Dict:
        """Main extraction method that routes to appropriate extractor"""
        
        # Validate file
        if not self.is_supported_file(filename):
            return {
                'success': False,
                'error': 'نوع الملف غير مدعوم. يرجى رفع ملف PDF أو Word',
                'error_type': 'unsupported_file_type'
            }
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if not self.validate_file_size(file_size):
            return {
                'success': False,
                'error': f'حجم الملف كبير جداً. الحد الأقصى {self.max_file_size // (1024*1024)} ميجابايت',
                'error_type': 'file_too_large'
            }
        
        # Extract based on file type
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        else:
            return {
                'success': False,
                'error': 'نوع الملف غير مدعوم',
                'error_type': 'unsupported_file_type'
            }
    
    def get_file_info(self, file_path: str, filename: str) -> Dict:
        """Get basic file information without extracting content"""
        try:
            file_size = os.path.getsize(file_path)
            ext = os.path.splitext(filename)[1].lower()
            
            return {
                'filename': filename,
                'size_bytes': file_size,
                'size_mb': round(file_size / (1024 * 1024), 2),
                'extension': ext,
                'is_supported': self.is_supported_file(filename),
                'size_valid': self.validate_file_size(file_size)
            }
        except Exception as e:
            return {
                'error': f'خطأ في قراءة معلومات الملف: {str(e)}'
            }

